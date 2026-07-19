"""Final-output rendering — convert the working manuscript (markdown) into the user's chosen
format: DOCX (template-aware), PDF, HTML, or Markdown.

DOCX inherits the target journal's uploaded .docx template (fonts / margins / heading styles)
when one exists, so output drops straight into the journal's format (e.g. the JEI template).
PDF is rendered with reportlab; HTML/Markdown reuse the editor's renderer. Files are written
under generated/<fmt>/ and returned to the caller.
"""
from __future__ import annotations

import re
from pathlib import Path

from .config import GENERATED, UPLOADS
from .agents.editor import md_to_html

FORMATS = ["docx", "pdf", "html", "markdown"]
_EXT = {"docx": "docx", "pdf": "pdf", "html": "html", "markdown": "md"}
_MIME = {"docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
         "pdf": "application/pdf", "html": "text/html", "markdown": "text/markdown"}


def mime(fmt: str) -> str:
    return _MIME.get(fmt, "application/octet-stream")


def _safe(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_") or "manuscript"


def _blocks(md: str):
    """Yield (kind, level, text) blocks: ('h', n, text) or ('p', 0, text)."""
    para: list[str] = []
    for ln in md.splitlines():
        h = re.match(r"^(#{1,6})\s+(.*)$", ln)
        if h:
            if para:
                yield ("p", 0, " ".join(para)); para = []
            yield ("h", len(h.group(1)), h.group(2).strip())
        elif ln.strip() == "":
            if para:
                yield ("p", 0, " ".join(para)); para = []
        else:
            para.append(ln.strip())
    if para:
        yield ("p", 0, " ".join(para))


def _journal_template_docx(journal: str) -> Path | None:
    if not journal:
        return None
    d = UPLOADS / "journals" / journal
    if d.exists():
        for f in sorted(d.glob("*.docx")):
            return f
    return None


# ── DOCX (template-aware) ─────────────────────────────────────────────────────
def to_docx(md: str, title: str, journal: str, profile: dict, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, Inches

    template = _journal_template_docx(journal)
    if template:
        doc = Document(str(template))
        body = doc.element.body
        for child in list(body):                       # blank the template body, keep page setup
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)
    else:
        doc = Document()
        fmt = (profile or {}).get("formatting") or {}
        try:
            m = float(re.sub(r"[^0-9.]", "", str(fmt.get("margins") or "1")) or 1)
            for s in doc.sections:
                s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(m)
        except Exception:  # noqa: BLE001
            pass
        normal = doc.styles["Normal"]
        normal.font.name = re.sub(r"[^A-Za-z ].*", "", str(fmt.get("font") or "Times New Roman")).strip() or "Times New Roman"
        try:
            normal.font.size = Pt(float(re.sub(r"[^0-9.]", "", str(fmt.get("font_size") or "12")) or 12))
        except Exception:  # noqa: BLE001
            normal.font.size = Pt(12)

    styles = {s.name for s in doc.styles}
    if title and not md.lstrip().startswith("#"):     # avoid duplicate title if md already has one
        _add_heading(doc, title, 0, styles)
    for kind, lvl, text in _blocks(md):
        if kind == "h":
            _add_heading(doc, text, min(max(lvl, 1), 4), styles)
        else:
            _add_runs(doc.add_paragraph(), text)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _add_heading(doc, text: str, level: int, styles: set) -> None:
    """Use the template's built-in Heading style when present; otherwise format a normal
    paragraph manually (bold + size) so headings survive templates like JEI that omit them."""
    from docx.shared import Pt

    style_name = "Title" if level == 0 else f"Heading {level}"
    if style_name in styles:
        try:
            doc.add_heading(text, level=level)
            return
        except Exception:  # noqa: BLE001
            pass
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt({0: 18, 1: 15, 2: 13, 3: 12, 4: 11}.get(level, 12))


def _add_runs(paragraph, text: str) -> None:
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        run = paragraph.add_run(part)
        run.bold = bool(i % 2)


# ── PDF (reportlab) ───────────────────────────────────────────────────────────
def to_pdf(md: str, title: str, profile: dict, out_path: Path) -> Path:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    out_path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName="Helvetica", fontSize=11,
                          leading=16, spaceAfter=8)
    doc = SimpleDocTemplate(str(out_path), pagesize=letter, topMargin=inch, bottomMargin=inch,
                            leftMargin=inch, rightMargin=inch, title=title)
    flow = []
    if title and not md.lstrip().startswith("#"):
        flow.append(Paragraph(_esc(title), styles["Title"]))
        flow.append(Spacer(1, 10))
    for kind, lvl, text in _blocks(md):
        if kind == "h":
            flow.append(Paragraph(_esc(text), styles[f"Heading{min(max(lvl,1),4)}"]))
        else:
            flow.append(Paragraph(_bold(_esc(text)), body))
    doc.build(flow)
    return out_path


def _esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _bold(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)


# ── dispatch ──────────────────────────────────────────────────────────────────
def render(md: str, *, title: str, journal: str, profile: dict, fmt: str) -> tuple[Path, str]:
    """Render the manuscript to `fmt`, write it under generated/<fmt>/, return (path, mime)."""
    fmt = fmt.lower()
    if fmt not in FORMATS:
        raise ValueError(f"unsupported format '{fmt}' (allowed: {FORMATS})")
    out = GENERATED / fmt / f"{_safe(title)}.{_EXT[fmt]}"
    out.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "markdown":
        out.write_text(md)
    elif fmt == "html":
        out.write_text(md_to_html(md, title))
    elif fmt == "docx":
        to_docx(md, title, journal, profile, out)
    elif fmt == "pdf":
        to_pdf(md, title, profile, out)
    return out, mime(fmt)
