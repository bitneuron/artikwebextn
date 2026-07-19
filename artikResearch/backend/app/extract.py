"""Document text extraction — PDF (PyMuPDF→pypdf fallback), DOCX, LaTeX, Markdown, plain text.

Returns plain text; the Paper Reader / Journal Reader agents do the structured understanding.
Never raises on a bad file — returns a best-effort string + a note.
"""
from __future__ import annotations

import re
from pathlib import Path


def detect_format(filename: str) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    return {"pdf": "pdf", "docx": "docx", "doc": "docx", "tex": "latex", "latex": "latex",
            "md": "markdown", "markdown": "markdown", "txt": "text", "html": "html",
            "htm": "html"}.get(ext, ext or "text")


def extract_text(path: str, fmt: str | None = None) -> str:
    p = Path(path)
    if not p.exists():
        return ""
    fmt = fmt or detect_format(p.name)
    try:
        if fmt == "pdf":
            return _pdf(p)
        if fmt == "docx":
            return _docx(p)
        if fmt == "html":
            return _html(p)
        # latex / markdown / text — read raw
        return p.read_text(errors="replace")
    except Exception as e:  # noqa: BLE001
        return f"(extraction error: {e})"


def _pdf(p: Path) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(p))
        return "\n\n".join(page.get_text() for page in doc)
    except Exception:  # noqa: BLE001
        try:
            from pypdf import PdfReader
            r = PdfReader(str(p))
            return "\n\n".join((pg.extract_text() or "") for pg in r.pages)
        except Exception as e:  # noqa: BLE001
            return f"(pdf extraction failed: {e})"


def _docx(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    parts = [para.text for para in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _html(p: Path) -> str:
    raw = p.read_text(errors="replace")
    raw = re.sub(r"<script.*?</script>", " ", raw, flags=re.S | re.I)
    raw = re.sub(r"<style.*?</style>", " ", raw, flags=re.S | re.I)
    return re.sub(r"<[^>]+>", " ", raw)


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text or ""))


def clip(text: str, max_chars: int = 24000) -> str:
    """Clip long docs for the LLM context window (keeps head + tail — abstract + refs)."""
    text = text or ""
    if len(text) <= max_chars:
        return text
    head = text[: int(max_chars * 0.7)]
    tail = text[-int(max_chars * 0.3):]
    return head + "\n\n…[middle truncated]…\n\n" + tail
